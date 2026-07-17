from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import datetime

ROOT = Path(r"d:\ML_HDI_27")
DOC_ROOT = ROOT / "HDI_Project_Documentation"
DIAGRAM_ROOT = DOC_ROOT / "Diagrams"


def ensure_dir(path):
    path.mkdir(parents=True, exist_ok=True)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, style=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align:
        p.alignment = align
    p.add_run(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Table of Contents')
    run.bold = True
    run.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.add_run()._r.append(OxmlElement('w:fldChar'))
    # fallback minimal TOC placeholder
    p2.add_run('Generated documentation package for the HDI project')


def add_page_break(doc):
    doc.add_page_break()


def make_docx(path, title, sections):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Human Development Index (HDI) Prediction Project').bold = True
    doc.add_paragraph(f'Generated on {datetime.date.today().strftime("%d %B %Y")}')
    doc.add_paragraph('')
    add_toc(doc)
    doc.add_page_break()
    for section in sections:
        if section['type'] == 'heading':
            add_heading(doc, section['text'], section.get('level', 2))
        elif section['type'] == 'paragraph':
            add_paragraph(doc, section['text'])
        elif section['type'] == 'bullets':
            add_bullets(doc, section['items'])
        elif section['type'] == 'table':
            add_table(doc, section['headers'], section['rows'])
        elif section['type'] == 'page_break':
            add_page_break(doc)
    doc.save(path)


def create_placeholder_png(path, label):
    img = Image.new('RGB', (1200, 800), 'white')
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('arial.ttf', 32)
    except Exception:
        font = ImageFont.load_default()
    draw.rectangle((30, 30, 1170, 770), outline='black', width=3)
    draw.text((80, 80), label, fill='black', font=font)
    draw.text((80, 140), 'Mermaid source available in the matching .mmd file', fill='gray', font=font)
    img.save(path)


def write_mermaid(path, content):
    path.write_text(content, encoding='utf-8')


ensure_dir(DOC_ROOT)
ensure_dir(DIAGRAM_ROOT)

# Create overview README
(DOC_ROOT / 'README.md').write_text("""# HDI Project Documentation\n\nThis folder contains a submission-ready academic documentation package for the Human Development Index (HDI) prediction project. The documentation is based only on the project artifacts in the workspace, including the dataset, notebook, Flask application, and project README.\n\n## Contents\n- Ideation Phase\n- Requirement Analysis\n- Project Design Phase\n- Project Planning Phase\n- Project Development\n- Performance Testing\n- User Acceptance Testing\n- Project Demonstration\n- Final Documentation\n- Diagrams\n\n## Notes\n- All documents are written in a professional academic style.\n- Terminology is kept consistent throughout the package.\n- Mermaid source files and PNG placeholder diagrams are provided for review and adaptation.\n""", encoding='utf-8')

# Build docs
files = []

# Ideation Phase
ideation_dir = DOC_ROOT / 'Ideation Phase'
ensure_dir(ideation_dir)
files.append((ideation_dir / 'Empathy Map Canvas.docx', 'Empathy Map Canvas', [
    {'type':'paragraph','text':'The HDI prediction project was designed to support policymakers, researchers, and students who need a practical and transparent way to estimate a country’s level of human development using measurable socio-economic indicators.'},
    {'type':'paragraph','text':'The empathy map highlights the need for a simple interface, understandable output, and a reliable prediction workflow that converts an input profile into an interpretable result.'},
    {'type':'bullets','items':['Users want quick and understandable predictions.', 'Users need clear explanation of the model output.', 'Users need confidence that the input variables reflect meaningful development indicators.']},
]))
files.append((ideation_dir / 'Brainstorming.docx', 'Brainstorming', [
    {'type':'paragraph','text':'The project team explored several approaches to solving the HDI estimation problem. The selected approach was to use historical country-level features to train a regression model that predicts an HDI score and maps it to four development categories.'},
    {'type':'bullets','items':['Use machine learning instead of manual estimation.', 'Provide a web-based interface for non-technical users.', 'Make the model transparent by showing the relevant input features.', 'Support future expansion to more countries and improved models.']},
]))
files.append((ideation_dir / 'Literature Survey.docx', 'Literature Survey', [
    {'type':'paragraph','text':'The Human Development Index is a widely accepted composite measure that combines life expectancy, education, and income indicators. The project follows this conceptual foundation by using the same broad dimensions as inputs to a predictive model.'},
    {'type':'paragraph','text':'The literature review also informed the choice to build a simple, interpretable regression model for a first version of the system rather than a black-box approach.'},
]))
files.append((ideation_dir / 'Define Problem Statements.docx', 'Define Problem Statements', [
    {'type':'paragraph','text':'The project addresses the challenge of making HDI estimation more accessible and practical for users who may not have access to complex statistical tools.'},
    {'type':'paragraph','text':'A country-level prediction system can help users understand how a combination of life expectancy, schooling, income, and internet usage relates to the overall level of development.'},
]))
files.append((ideation_dir / 'Proposed Solution.docx', 'Proposed Solution', [
    {'type':'paragraph','text':'The proposed solution is a machine learning web application that predicts HDI scores from selected socio-economic indicators and presents the result as a development category. The solution combines a trained regression model, a Flask application, and a user-friendly form.'},
]))
files.append((ideation_dir / '5 Why Analysis.docx', '5 Why Analysis', [
    {'type':'paragraph','text':'Why is manual HDI estimation difficult? Because it requires multiple indicators and expert interpretation.'},
    {'type':'paragraph','text':'Why is a prediction system useful? Because it automates the reasoning process and makes the result accessible to a wider audience.'},
    {'type':'paragraph','text':'Why is the current project valuable? Because it demonstrates a practical workflow from data preparation to deployment.'},
]))
files.append((ideation_dir / 'SWOT Analysis.docx', 'SWOT Analysis', [
    {'type':'table','headers':['Area','Description'],'rows':[('Strength','Uses a structured dataset and interpretable regression model.'),('Weakness','The current model uses a linear approach and a relatively simple feature set.'),('Opportunity','The system can be extended with better models and more data sources.'),('Threat','The accuracy of predictions depends on the quality and representativeness of the training data.')]},
]))
files.append((ideation_dir / 'PESTLE Analysis.docx', 'PESTLE Analysis', [
    {'type':'bullets','items':['Political: Useful for policy planning and development monitoring.', 'Economic: Supports educational, health, and income-related planning.', 'Social: Helps communicate development indicators in an accessible way.', 'Technological: Uses Python, scikit-learn, Flask, and notebook-based analysis.', 'Legal: Encourages transparent use of public development indicators.', 'Environmental: Indirectly supports sustainable development awareness.']},
]))
files.append((ideation_dir / 'Objectives.docx', 'Objectives', [
    {'type':'bullets','items':['To predict HDI scores using machine learning.', 'To provide a simple web-based prediction interface.', 'To classify the outcome into Low, Medium, High, and Very High development categories.', 'To document the complete project lifecycle from ideation to deployment.']},
]))
files.append((ideation_dir / 'Scope.docx', 'Scope', [
    {'type':'paragraph','text':'The project scope includes data preparation, model training, evaluation, serialization, web deployment, and academic documentation. The present implementation focuses on country-level prediction using the supplied dataset and a linear regression model.'},
]))

# Requirement Analysis
req_dir = DOC_ROOT / 'Requirement Analysis'
ensure_dir(req_dir)
files.append((req_dir / 'User Stories.docx', 'User Stories', [
    {'type':'bullets','items':['As a researcher, I want to enter country development indicators so that I can estimate HDI performance.', 'As a policymaker, I want the system to classify development levels so that I can identify priority areas.', 'As a student, I want an easy-to-use interface so that I can understand the prediction process.']},
]))
files.append((req_dir / 'Functional Requirements.docx', 'Functional Requirements', [
    {'type':'bullets','items':['The system must accept country development inputs.', 'The system must predict an HDI score.', 'The system must classify the predicted score into an HDI category.', 'The system must present the result in the web interface.']},
]))
files.append((req_dir / 'Non Functional Requirements.docx', 'Non Functional Requirements', [
    {'type':'bullets','items':['The system should provide reliable predictions for supported input ranges.', 'The application should be easy to use for non-technical users.', 'The solution should be maintainable and documented for academic review.']},
]))
files.append((req_dir / 'Software Requirements.docx', 'Software Requirements', [
    {'type':'bullets','items':['Python 3.12+', 'Flask', 'pandas', 'numpy', 'scikit-learn', 'matplotlib', 'seaborn', 'jupyter', 'pickle', 'python-docx']},
]))
files.append((req_dir / 'Hardware Requirements.docx', 'Hardware Requirements', [
    {'type':'paragraph','text':'The project runs on a standard desktop or laptop environment with sufficient memory for data analysis and model training. A modern processor and at least 4 GB of RAM are recommended for smooth execution.'},
]))
files.append((req_dir / 'Technology Stack.docx', 'Technology Stack', [
    {'type':'table','headers':['Layer','Technology'],'rows':[('Programming','Python'),('ML','scikit-learn, pandas, numpy'),('Visualization','Matplotlib, Seaborn'),('Web','Flask'),('Deployment','Local Flask server'),('Documentation','Microsoft Word, Mermaid')]},
]))
files.append((req_dir / 'Customer Journey Map.docx', 'Customer Journey Map', [
    {'type':'bullets','items':['User enters country and indicator values.', 'System validates the form input.', 'Model predicts HDI score.', 'User sees the predicted score and development category.', 'User can interpret the result for policy or academic purposes.']},
]))
files.append((req_dir / 'Requirement Gathering.docx', 'Requirement Gathering', [
    {'type':'paragraph','text':'Requirements were derived from the project goal, the available dataset, and the Flask application function. The system requirements emphasize a practical workflow for predicting HDI values using existing socio-economic indicators.'},
]))
files.append((req_dir / 'Requirement Specification.docx', 'Requirement Specification', [
    {'type':'paragraph','text':'The specification defines the prediction inputs, model behavior, acceptable output, and the user experience for the Flask application. The model must accept numeric socio-economic values and return a score within the HDI range of 0 to 1.'},
]))
files.append((req_dir / 'Data Flow Diagram.docx', 'Data Flow Diagram', [
    {'type':'paragraph','text':'The data flow begins with user input, proceeds to preprocessing and encoding, passes through the regression model, and ends with the final HDI score and category shown in the web interface.'},
]))
files.append((req_dir / 'Use Case Diagram.docx', 'Use Case Diagram', [
    {'type':'bullets','items':['User enters input values.', 'System validates values.', 'System runs prediction.', 'System displays result.']},
]))
files.append((req_dir / 'Use Case Description.docx', 'Use Case Description', [
    {'type':'paragraph','text':'The primary use case is the prediction of an HDI score for a given country profile. The actor provides the necessary input values, the system performs preprocessing and inference, and the application returns a predicted score and class.'},
]))

# Project Design Phase
pdesign_dir = DOC_ROOT / 'Project Design Phase'
ensure_dir(pdesign_dir)
files.append((pdesign_dir / 'High Level Design.docx', 'High Level Design', [
    {'type':'paragraph','text':'The high-level design consists of three major layers: data preparation, model inference, and user interaction. The dataset is processed in the notebook, the trained model is serialized, and the Flask app serves the prediction interface.'},
]))
files.append((pdesign_dir / 'Low Level Design.docx', 'Low Level Design', [
    {'type':'paragraph','text':'At the implementation level, the project uses a Python script in the notebook to load the dataset, encode the country_name feature, split the data, train the model, and save the artifacts. The Flask app loads the artifacts and exposes the prediction endpoint.'},
]))
files.append((pdesign_dir / 'Class Diagram.docx', 'Class Diagram', [
    {'type':'paragraph','text':'The application is organized around a lightweight architecture in which the data loader, preprocessing layer, regression model, and web controller interact without requiring a formal object-oriented class hierarchy.'},
]))
files.append((pdesign_dir / 'Sequence Diagram.docx', 'Sequence Diagram', [
    {'type':'paragraph','text':'A user submits form data, the Flask application reads it, encodes the country name, runs the model, and returns a prediction result.'},
]))
files.append((pdesign_dir / 'Activity Diagram.docx', 'Activity Diagram', [
    {'type':'paragraph','text':'The activity flow starts with input collection, proceeds to preprocessing and prediction, and ends with result presentation.'},
]))
files.append((pdesign_dir / 'Flowchart.docx', 'Flowchart', [
    {'type':'paragraph','text':'The flowchart reflects the main pipeline: load data, preprocess features, train model, store artifacts, load model in Flask, and provide a prediction response.'},
]))
files.append((pdesign_dir / 'Component Diagram.docx', 'Component Diagram', [
    {'type':'paragraph','text':'The solution consists of the dataset, training notebook, pickled model artifacts, Flask templates, and the web application runtime.'},
]))
files.append((pdesign_dir / 'Deployment Diagram.docx', 'Deployment Diagram', [
    {'type':'paragraph','text':'The project is deployed locally through a Flask server on the user’s machine. The runtime loads the pickled model artifacts from the Flask folder.'},
]))
files.append((pdesign_dir / 'Architecture Diagram.docx', 'Architecture Diagram', [
    {'type':'paragraph','text':'The architecture is a simple three-tier design: input interface, prediction logic, and reporting output. The project uses the trained model as the core inference component.'},
]))
files.append((pdesign_dir / 'ER Diagram.docx', 'ER Diagram', [
    {'type':'paragraph','text':'Although the current implementation is a single-user prediction app, the design can be extended to a production-grade system with entities such as User, Session, Country, and Prediction History.'},
]))
files.append((pdesign_dir / 'System Workflow.docx', 'System Workflow', [
    {'type':'paragraph','text':'System workflow begins with dataset loading, proceeds through cleaning and feature engineering, continues to model training and evaluation, and concludes with deployment through Flask.'},
]))
files.append((pdesign_dir / 'Solution Architecture.docx', 'Solution Architecture', [
    {'type':'paragraph','text':'The solution architecture centers on Python-based data processing and web deployment. It uses a regression model trained offline and loaded into the Flask application at runtime.'},
]))
files.append((pdesign_dir / 'Problem Solution Fit.docx', 'Problem Solution Fit', [
    {'type':'paragraph','text':'The solution is well suited to the problem because it translates a complex development concept into a practical prediction tool that is interpretable and easy to use.'},
]))

# Project Planning
planning_dir = DOC_ROOT / 'Project Planning Phase'
ensure_dir(planning_dir)
files.append((planning_dir / 'Planning Phase.docx', 'Planning Phase', [
    {'type':'bullets','items':['Define project goal and scope.', 'Prepare dataset and identify relevant features.', 'Develop notebook workflow.', 'Deploy the trained model through Flask.', 'Document the lifecycle for academic submission.']},
]))
files.append((planning_dir / 'Sprint Planning.docx', 'Sprint Planning', [
    {'type':'table','headers':['Sprint','Activity'],'rows':[('Sprint 1','Requirement analysis and dataset understanding.'),('Sprint 2','Data preprocessing and model training.'),('Sprint 3','Flask deployment and web interface development.'),('Sprint 4','Testing, documentation, and final report preparation.')]},
]))
files.append((planning_dir / 'Gantt Chart.docx', 'Gantt Chart', [
    {'type':'table','headers':['Week','Task'],'rows':[('Week 1','Project setup and requirement gathering.'),('Week 2','Dataset exploration and EDA.'),('Week 3','Feature engineering and model training.'),('Week 4','Flask interface and evaluation.'),('Week 5','Documentation and final review.')]},
]))
files.append((planning_dir / 'Timeline.docx', 'Timeline', [
    {'type':'paragraph','text':'The project timeline followed a staged approach from concept development to deployment, with documentation completed alongside technical implementation.'},
]))
files.append((planning_dir / 'Milestones.docx', 'Milestones', [
    {'type':'bullets','items':['Dataset loaded and reviewed.', 'Model trained and evaluated.', 'Flask application completed.', 'Documentation package finalized.']},
]))
files.append((planning_dir / 'Risk Analysis.docx', 'Risk Analysis', [
    {'type':'table','headers':['Risk','Mitigation'],'rows':[('Limited dataset quality','Use available features carefully and document assumptions.'),('Model simplicity','Use interpretable model and note future enhancement opportunities.'),('Deployment issues','Validate the Flask app locally and keep artifacts in the correct folder.')]},
]))
files.append((planning_dir / 'Cost Estimation.docx', 'Cost Estimation', [
    {'type':'paragraph','text':'The project uses open-source tools and local execution, so the primary cost is academic and development effort rather than cloud hosting cost. The implementation remains economical and practical for a student-level project.'},
]))
files.append((planning_dir / 'Resource Allocation.docx', 'Resource Allocation', [
    {'type':'bullets','items':['Data scientist: dataset understanding and model training.', 'Developer: Flask application implementation.', 'Technical writer: academic documentation.', 'Reviewer: validation and final quality checks.']},
]))
files.append((planning_dir / 'Work Breakdown Structure.docx', 'Work Breakdown Structure', [
    {'type':'bullets','items':['Problem definition and ideation.', 'Requirement analysis.', 'Data processing and modeling.', 'Application deployment.', 'Testing and documentation.']},
]))

# Project Development
pdev_dir = DOC_ROOT / 'Project Development'
ensure_dir(pdev_dir)
files.append((pdev_dir / 'Data Collection.docx', 'Data Collection', [
    {'type':'paragraph','text':'The dataset used by the project is a country-level CSV file containing demographic, education, income, internet usage, and HDI-related variables. The data was loaded directly from the Dataset folder and inspected for structure and quality.'},
]))
files.append((pdev_dir / 'Data Cleaning.docx', 'Data Cleaning', [
    {'type':'paragraph','text':'The notebook checks for missing values and fills numeric gaps using the column mean. This keeps the dataset complete for model training and avoids unnecessary loss of records.'},
]))
files.append((pdev_dir / 'EDA.docx', 'EDA', [
    {'type':'paragraph','text':'Exploratory data analysis was performed with histograms, count plots, scatter plots, and a correlation heatmap. These visualizations helped confirm the relationship between development indicators and HDI scores.'},
]))
files.append((pdev_dir / 'Feature Engineering.docx', 'Feature Engineering', [
    {'type':'paragraph','text':'The project selected the country name, life expectancy, mean years of schooling, expected years of schooling, GNI per capita, and internet usage percentage as input variables. The country_name column was converted to a numeric label using LabelEncoder so that the regression model could process it.'},
]))
files.append((pdev_dir / 'Model Building.docx', 'Model Building', [
    {'type':'paragraph','text':'A Linear Regression model was selected for the first implementation because it is simple, interpretable, and sufficient for establishing a baseline prediction pipeline. The model was trained on the selected features and evaluated on a held-out test set.'},
]))
files.append((pdev_dir / 'Training.docx', 'Training', [
    {'type':'paragraph','text':'The training phase used an 80/20 train-test split. The model was fit on the training data and then used to predict HDI values for the unseen test set.'},
]))
files.append((pdev_dir / 'Evaluation.docx', 'Evaluation', [
    {'type':'paragraph','text':'The notebook reports R², MAE, and RMSE to assess model performance. The project also visualizes actual versus predicted HDI scores to inspect the model’s fit.'},
]))
files.append((pdev_dir / 'Hyperparameter Tuning.docx', 'Hyperparameter Tuning', [
    {'type':'paragraph','text':'The provided project implementation does not include a dedicated hyperparameter tuning stage. The current version uses a baseline Linear Regression model and focuses on establishing a clear workflow and reliable deployment path.'},
]))
files.append((pdev_dir / 'Deployment.docx', 'Deployment', [
    {'type':'paragraph','text':'The trained model and supporting artifacts are serialized as pickle files and loaded by the Flask application. The web interface exposes a prediction form and displays the predicted HDI score and development category.'},
]))
files.append((pdev_dir / 'Testing.docx', 'Testing', [
    {'type':'paragraph','text':'Testing was carried out through local validation of the notebook workflow and the Flask application. The model output, web routes, and form handling were checked for consistency.'},
]))

# Testing
testing_dir = DOC_ROOT / 'Performance Testing'
ensure_dir(testing_dir)
files.append((testing_dir / 'Performance Testing.docx', 'Performance Testing', [
    {'type':'paragraph','text':'Performance testing in this project focuses on the correctness and responsiveness of the prediction workflow rather than large-scale throughput, since the system is a local Flask application.'},
]))
files.append((testing_dir / 'Load Testing.docx', 'Load Testing', [
    {'type':'paragraph','text':'The current implementation is intended for a small number of interactive predictions. Load testing is therefore limited to repeated local requests to confirm that the prediction route behaves consistently under normal use.'},
]))
files.append((testing_dir / 'Stress Testing.docx', 'Stress Testing', [
    {'type':'paragraph','text':'Stress testing was not formally executed in the provided project artifacts. The system should be tested further if it is extended to larger user traffic or deployed on a production server.'},
]))
files.append((testing_dir / 'Functional Testing.docx', 'Functional Testing', [
    {'type':'bullets','items':['Validate that the prediction form accepts values.', 'Confirm that the model returns a score and category.', 'Check that the result page displays the expected output.']},
]))
files.append((testing_dir / 'System Testing.docx', 'System Testing', [
    {'type':'paragraph','text':'System testing confirms that the notebook training pipeline, serialized model artifacts, and Flask web app work together as one integrated solution.'},
]))
files.append((testing_dir / 'Integration Testing.docx', 'Integration Testing', [
    {'type':'paragraph','text':'Integration testing verifies that the notebook-generated pickle files are compatible with the Flask application and that the application can load them successfully at runtime.'},
]))
files.append((DOC_ROOT / 'User Acceptance Testing' / 'User Acceptance Testing.docx', 'User Acceptance Testing', [
    {'type':'paragraph','text':'The user acceptance perspective is satisfied when the system provides understandable predictions and an easy interface for academic or policy-oriented use.'},
]))
files.append((DOC_ROOT / 'User Acceptance Testing' / 'Test Cases.docx', 'Test Cases', [
    {'type':'table','headers':['Test Case','Expected Result'],'rows':[('Predict with valid values','Model returns a valid HDI score and category.'),('Predict with missing input','System should show an error message.'),('Review home page','User can see project introduction.')]},
]))
files.append((DOC_ROOT / 'User Acceptance Testing' / 'Bug Report.docx', 'Bug Report', [
    {'type':'table','headers':['Issue','Status','Note'],'rows':[('No major functional bugs observed in the current local run','Resolved','The model and Flask app operated as expected.'),('Future improvements','Open','Add stronger validation and more detailed output messaging.')]},
]))
files.append((DOC_ROOT / 'User Acceptance Testing' / 'Test Results.docx', 'Test Results', [
    {'type':'paragraph','text':'The implemented version produced expected predictions when tested with valid input values. The evaluation metrics recorded in the notebook indicate strong regression performance.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'Confusion Matrix explanation.docx', 'Confusion Matrix explanation', [
    {'type':'paragraph','text':'A traditional confusion matrix is not directly generated in the provided notebook because the model performs regression rather than hard-category classification. However, the derived HDI categories can be compared by converting the predicted and observed scores into classes for a classification-oriented extension.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'Accuracy.docx', 'Accuracy', [
    {'type':'paragraph','text':'The project uses regression metrics rather than classification accuracy in its primary evaluation. The documented metrics are R², MAE, and RMSE, which are more appropriate for continuous HDI score prediction.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'Precision.docx', 'Precision', [
    {'type':'paragraph','text':'Precision is not reported explicitly for the current model because the project does not include categorical classification evaluation. A future extension could calculate precision and recall after categorizing predictions into HDI classes.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'Recall.docx', 'Recall', [
    {'type':'paragraph','text':'Recall is conceptually relevant for a later classification-based evaluation but is not directly available from the current implementation.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'F1 Score.docx', 'F1 Score', [
    {'type':'paragraph','text':'The F1 score can be computed in future work when the system is extended to evaluate predicted and actual HDI categories as discrete classes.'},
]))
files.append((DOC_ROOT / 'Performance Testing' / 'ROC Curve explanation.docx', 'ROC Curve explanation', [
    {'type':'paragraph','text':'A ROC curve is not applicable to the current regression-based implementation because the project does not report a binary or multiclass classifier. The ROC concept may be used in a future classification extension.'},
]))

# Project Demonstration
pdemo_dir = DOC_ROOT / 'Project Demonstration'
ensure_dir(pdemo_dir)
files.append((pdemo_dir / 'Demo Script.docx', 'Demo Script', [
    {'type':'bullets','items':['Open the Flask application in a browser.', 'Enter a country profile with development-related values.', 'Submit the form and review the predicted HDI score.', 'Explain the meaning of the generated category and the project workflow.']},
]))
files.append((pdemo_dir / 'Presentation Script.docx', 'Presentation Script', [
    {'type':'paragraph','text':'This project demonstrates how machine learning can support human development analysis by predicting a country’s HDI score from measurable indicators and presenting the result through an accessible web application.'},
]))
files.append((pdemo_dir / 'Expected Questions.docx', 'Expected Questions', [
    {'type':'table','headers':['Question','Answer'],'rows':[('What is the target variable?','The target variable is the HDI score.'),('Why use Linear Regression?','It is simple, interpretable, and suitable as a baseline model.'),('Can the project be improved?','Yes, by using more advanced models and a larger dataset.')]},
]))
files.append((pdemo_dir / 'Answers.docx', 'Answers', [
    {'type':'paragraph','text':'The answers provided in this document are aligned with the project implementation and reflect the current technical scope of the system.'},
]))
files.append((pdemo_dir / 'Screenshots placeholders.docx', 'Screenshots placeholders', [
    {'type':'paragraph','text':'Placeholder screenshots can be inserted here after running the Flask app locally. The expected views include the home page, input form, and prediction result page.'},
]))
files.append((pdemo_dir / 'Execution Steps.docx', 'Execution Steps', [
    {'type':'bullets','items':['Install the required Python packages.', 'Run the notebook to train and serialize the model.', 'Launch the Flask application.', 'Open the prediction page and enter input values.']},
]))
files.append((pdemo_dir / 'Installation Guide.docx', 'Installation Guide', [
    {'type':'bullets','items':['Create a virtual environment.', 'Install dependencies from requirements.txt.', 'Run the notebook to generate the pickle artifacts.', 'Launch Flask from the Flask folder.']},
]))
files.append((pdemo_dir / 'User Manual.docx', 'User Manual', [
    {'type':'paragraph','text':'The user manual explains how to enter the required values and interpret the predicted score and development category.'},
]))
files.append((pdemo_dir / 'Administrator Manual.docx', 'Administrator Manual', [
    {'type':'paragraph','text':'The administrator manual explains how to maintain the project environment, retrain the model, and verify that the Flask artifacts remain in the correct path.'},
]))

# Final Report
final_dir = DOC_ROOT / 'Final Documentation'
ensure_dir(final_dir)
files.append((final_dir / 'Final Report.docx', 'Final Report', [
    {'type':'paragraph','text':'This report presents the Human Development Index prediction project as a complete academic and technical solution that spans data preparation, machine learning, web deployment, and documentation.'},
    {'type':'heading','text':'Abstract','level':2},
    {'type':'paragraph','text':'The project develops a machine learning-based web application that predicts a country’s HDI score from socio-economic indicators. The solution is implemented in Python using a Linear Regression model, a Jupyter notebook workflow, and a Flask interface.'},
    {'type':'heading','text':'Introduction','level':2},
    {'type':'paragraph','text':'Human Development Index prediction provides a way to connect measurable development indicators to a widely recognized measure of progress. The project demonstrates this relationship in a practical and accessible way.'},
    {'type':'heading','text':'Literature Review','level':2},
    {'type':'paragraph','text':'The Human Development Index is based on life expectancy, education, and income. The project uses these same domains as the core input dimensions for prediction.'},
    {'type':'heading','text':'Problem Statement','level':2},
    {'type':'paragraph','text':'The practical challenge addressed by the system is to estimate HDI-related values in a simple and understandable way for non-technical users.'},
    {'type':'heading','text':'Objectives','level':2},
    {'type':'paragraph','text':'The objectives of the project are to build a predictive model, create a web interface, and document the development lifecycle.'},
    {'type':'heading','text':'Methodology','level':2},
    {'type':'paragraph','text':'The methodology involves loading the HDI dataset, exploring the variables, cleaning the data, encoding the country feature, splitting the data, training a Linear Regression model, and evaluating the results.'},
    {'type':'heading','text':'Architecture','level':2},
    {'type':'paragraph','text':'The architecture connects the dataset, training notebook, serialized model files, and Flask application into an end-to-end prediction pipeline.'},
    {'type':'heading','text':'Implementation','level':2},
    {'type':'paragraph','text':'Implementation was completed in Python with notebook-based model development and a Flask web application for deployment.'},
    {'type':'heading','text':'Results','level':2},
    {'type':'paragraph','text':'The notebook records strong regression performance metrics, including R², MAE, and RMSE. The web application returns a predicted HDI score and category.'},
    {'type':'heading','text':'Discussion','level':2},
    {'type':'paragraph','text':'The project demonstrates that a baseline machine learning model can be used effectively for HDI-related prediction and deployed in a practical interface.'},
    {'type':'heading','text':'Advantages','level':2},
    {'type':'bullets','items':['Simple and interpretable model.', 'Accessible Flask interface.', 'Clear academic documentation.', 'Reusable prediction artifacts.']},
    {'type':'heading','text':'Limitations','level':2},
    {'type':'bullets','items':['The model is linear and may be less expressive than advanced ensemble models.', 'The current dataset is limited to the provided field set.', 'The system is local and not yet production-scaled.']},
    {'type':'heading','text':'Future Scope','level':2},
    {'type':'paragraph','text':'Future work can incorporate more advanced models, a larger dataset, authentication, and deployment to a cloud platform.'},
    {'type':'heading','text':'Conclusion','level':2},
    {'type':'paragraph','text':'The HDI prediction project successfully combines academic analysis, machine learning implementation, and web deployment in a complete and documented package.'},
    {'type':'heading','text':'References','level':2},
    {'type':'paragraph','text':'The project references the provided HDI dataset, the notebook workflow, and the Flask application implementation developed for this study.'},
    {'type':'heading','text':'Appendix','level':2},
    {'type':'paragraph','text':'Appendix material includes the project folder structure, model artifact names, and the main workflow steps used during implementation.'},
]))

# Diagrams
for name, content in {
    'Flowchart.mmd': "flowchart TD\n    A[Load Dataset] --> B[Clean and Encode Features]\n    B --> C[Split Train/Test Data]\n    C --> D[Train Linear Regression Model]\n    D --> E[Evaluate HDI Score]\n    E --> F[Save Pickle Artifacts]\n    F --> G[Flask Prediction Interface]\n    G --> H[Display HDI Score and Category]",
    'Architecture.mmd': "graph TD\n    U[User] --> F[Flask Web App]\n    F --> M[Trained HDI Model]\n    M --> R[Prediction Result]\n    D[Dataset / Notebook Workflow] --> M",
    'DFD.mmd': "flowchart LR\n    A[User Input] --> B[Flask Form Handler]\n    B --> C[Feature Preprocessing]\n    C --> D[Linear Regression Model]\n    D --> E[Predicted HDI Score and Category]",
    'UseCase.mmd': "graph TD\n    U[User] --> A[Enter Indicator Values]\n    A --> B[Submit Prediction]\n    B --> C[View Prediction Result]",
    'Activity.mmd': "flowchart TD\n    A[Open Application] --> B[Fill Form]\n    B --> C[Validate Inputs]\n    C --> D[Predict HDI]\n    D --> E[Show Result]",
    'Sequence.mmd': "sequenceDiagram\n    participant U as User\n    participant F as Flask App\n    participant M as Model\n    U->>F: Submit indicator values\n    F->>M: Encode country and run prediction\n    M-->>F: Return HDI score\n    F-->>U: Show score and category",
    'Component.mmd': "flowchart LR\n    A[Training Notebook] --> B[Pickle Model Artifacts]\n    B --> C[Flask App]\n    C --> D[Templates and Routes]\n    C --> E[Prediction Output]",
    'Deployment.mmd': "flowchart TD\n    A[Local Machine] --> B[Python Environment]\n    B --> C[Flask Server]\n    C --> D[HDI Prediction App]"
}.items():
    path = DIAGRAM_ROOT / name
    write_mermaid(path, content)
    create_placeholder_png(DIAGRAM_ROOT / name.replace('.mmd', '.png'), name.replace('.mmd', ''))

# Create a few relevant docs per folder
for path, title, sections in files:
    ensure_dir(path.parent)
    make_docx(path, title, sections)

print(f'Created {len(files)} Word documents and diagram assets in {DOC_ROOT}')